from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A helper function to add a hyperlink to a paragraph object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

doc = Document()

# --- HEADER ---
name = doc.add_heading('BIRUT GUCHHAIT', 0)
name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

contact = doc.add_paragraph()
contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact.add_run('Gopalsingpur, West Bengal, 721434 | ‪+91-8170974624‬\n')
contact.add_run('Birutguchhait1@gmail.com | ')
add_hyperlink(contact, "http://www.linkedin.com/in/birut-guchhait-1206j2001", "LinkedIn")
contact.add_run(' | ')
add_hyperlink(contact, "https://github.com/BirutG", "GitHub")

# --- PROFESSIONAL SUMMARY ---
doc.add_heading('PROFESSIONAL SUMMARY', level=1)
doc.add_paragraph(
    "Results-oriented Data Analyst and MBA candidate with rigorous training in Business Analytics. "
    "Proficient in Python, SQL, and Power BI, with a proven track record of transforming raw data into actionable business insights. "
    "Demonstrated expertise in predictive modeling and customer segmentation through hands-on projects in banking and retail. "
    "Seeking to leverage skills in statistical analysis and data visualization to drive operational efficiency and revenue growth."
)

# --- TECHNICAL SKILLS ---
doc.add_heading('TECHNICAL SKILLS', level=1)
skills = [
    ("Programming & Scripting:", " Python (Pandas, NumPy, Matplotlib, Seaborn), R (Dplyr, Ggplot2), SQL (PostgreSQL, MySQL)."),
    ("Data Visualization & BI:", " Microsoft Power BI, Excel (Advanced Dashboards, Pivot Tables), KNIME."),
    ("Data Science Competencies:", " Exploratory Data Analysis (EDA), Feature Engineering, Predictive Modeling, PCA, Hypothesis Testing."),
    ("Soft Skills:", " Stakeholder Management, Problem Solving, Cross-functional Collaboration.")
]
for cat, desc in skills:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cat).bold = True
    p.add_run(desc)

# --- PROJECTS ---
doc.add_heading('PROJECTS', level=1)

# Project 1
p = doc.add_paragraph()
p.add_run("Customer Churn Analytics for Banking (SQL & Power BI)").bold = True
bullets = [
    "Developed complex SQL scripts using Common Table Expressions (CTEs) and Window Functions to analyze customer demographics and transactional behaviors.",
    "Identified a baseline churn rate of 20.03% and discovered that churned customers averaged 59.7 points lower in credit scores compared to retained peers.",
    "Segmented high-value 'at-risk' customers (top quartile balances), determining they were 1.8x more likely to churn, and recommended targeted retention incentives."
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# Project 2
p = doc.add_paragraph()
p.add_run("Shark Tank India Season 1 Analysis (Python)").bold = True
bullets = [
    "Performed comprehensive data cleaning and feature extraction on a dataset of 117 pitches using Python (Pandas) to analyze investment trends.",
    "Engineered metrics such as 'Deal Status' and 'Valuation,' revealing an average deal valuation of ₹44 Lakhs and identifying the most active investors.",
    "Visualized investment patterns and calculating shark participation rates to highlight successful pitch characteristics."
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# Project 3
p = doc.add_paragraph()
p.add_run("Wine Quality Prediction (KNIME)").bold = True
bullets = [
    "Built a machine learning predictive model to assess chemical dependencies (e.g., acidity, sugar) on wine quality ratings.",
    "Utilized the KNIME Analytics Platform for workflow automation and model validation."
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# Project 4
p = doc.add_paragraph()
p.add_run("Super Store Dataset Analysis (Excel)").bold = True
bullets = [
    "Designed dynamic dashboards to visualize regional sales trends, profit margins, and category performance.",
    "Derived actionable insights for inventory planning, identifying underperforming regions to optimize supply chain operations."
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# Project 5
p = doc.add_paragraph()
p.add_run("Principal Component Analysis (R)").bold = True
bullets = [
    "Conducted PCA on the Irish dataset to reduce dimensionality while retaining 95%+ of the variance.",
    "Identified key variables influencing class separation to improve the performance of subsequent classification models."
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# --- PROFESSIONAL EXPERIENCE ---
doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)

# Exp 1
p = doc.add_paragraph()
p.add_run("PS GROUP").bold = True
p.add_run(" | Data Analytics Intern")
doc.add_paragraph("September 2025 – November 2025").italic = True
bullets = [
    "Spearheaded the end-to-end collection and analysis of ESG (Environmental, Social, and Governance) data from multiple internal stakeholders.",
    "Constructed a centralized Excel database to track carbon emissions and sustainability metrics, reducing data retrieval time by 30% and facilitating data-driven decisions for green initiatives."
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# Exp 2
p = doc.add_paragraph()
p.add_run("MAHINDRA SWARAJ").bold = True
p.add_run(" | Sales & Marketing Intern")
doc.add_paragraph("March 2021 – April 2021").italic = True
bullets = [
    "Collaborated with cross-functional teams to execute region-specific promotional activities, contributing to the achievement of monthly sales targets.",
    "Analyzed customer feedback data to optimize local marketing strategies and improve customer engagement."
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# --- EDUCATION ---
doc.add_heading('EDUCATION', level=1)

# Education 1
p = doc.add_paragraph()
p.add_run("MBA (Business Analytics & Data Science)").bold = True
p.add_run(" | Raiganj University | Expected Aug 2026")

# Education 2
p = doc.add_paragraph()
p.add_run("BBA (Bachelor of Business Administration)").bold = True
p.add_run(" | MAKAUT | 2022")
doc.add_paragraph("Score: 8.97 CGPA (85.2%)", style='List Bullet')

# Save
doc.save('Birut_Guchhait_CV_Optimized.docx')
print("CV Generated Successfully!")